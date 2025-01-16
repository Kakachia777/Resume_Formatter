import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pdfminer.high_level import extract_text
import io
import logging
import os
import tempfile
import json
import re

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Remove or comment out the direct API key
# GEMINI_API_KEY = 'AIzaSyDdXpU3bn57KJEbA58rqjdK5yHccpOEbWs'

def get_gemini_api_key():
    """Get Gemini API key from environment or use default for development"""
    return os.getenv('GEMINI_API_KEY', 'AIzaSyDdXpU3bn57KJEbA58rqjdK5yHccpOEbWs')

# Gemini API setup with direct API key
genai.configure(api_key=get_gemini_api_key())
model = genai.GenerativeModel('gemini-2.0-flash-exp')

class ResumeProcessor:
    def __init__(self, doc: Document):
        self.doc = doc
        self.output_doc = Document()

    def extract_content(self, text: str) -> dict:
        prompt = f"""You are a resume parser. Extract information from the following resume text into a structured JSON format. Be thorough and extract all relevant information, even if fields appear empty at first glance.

Input Resume Text:
{text}

Instructions:
1. Extract ALL text that could represent the person's name, education, skills, and work experience
2. If certain sections are unclear, look for contextual clues
3. Include ALL technical skills mentioned anywhere in the resume
4. Capture complete work experience details including dates, titles, and bullet points
5. Do not leave fields empty - use available information from the entire text

Required JSON structure:
{{
    "name": "Full Name",
    "education": {{
        "university": "University Name",
        "degree": "Degree(s) Earned"
    }},
    "technical_skills": {{
        "Category1": ["Skill1", "Skill2"],
        "Category2": ["Skill3", "Skill4"],
        // Add more categories as needed
    }},
    "experience": [
        {{
            "company": "Company Name",
            "dates": "Start Date – End Date",
            "title": "Job Title",
            "bullets": [
                "Achievement/responsibility details"
            ]
        }}
    ]
}}

Output only the JSON object with all available information from the resume. No other text."""

        try:
            response = model.generate_content(prompt)
            if not response or not response.text:
                raise ValueError("Empty response from Gemini API")
                
            cleaned_response = response.text.strip()
            cleaned_response = cleaned_response.replace('```json', '').replace('```', '')
            
            logger.info(f"Extracted text content: {text[:500]}...")  # Log first 500 chars of input
            logger.info(f"Gemini response: {cleaned_response}")
            
            content = json.loads(cleaned_response)
            
            # Validate content has required fields
            if not content.get('name') or not content.get('education') or not content.get('experience'):
                raise ValueError("Missing required fields in parsed content")
            
            return content
            
        except json.JSONDecodeError as e:
            logger.error(f"JSON parsing error: {str(e)}")
            logger.error(f"Response text: {response.text}")
            raise
        except Exception as e:
            logger.error(f"Error in content extraction: {str(e)}")
            raise

    def add_header_image(self):
        section = self.output_doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = header_paragraph.add_run()
        run.add_picture('optomi_logo.png', width=Inches(3.2))

    def format_name(self, name: str):
        para = self.output_doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(name)
        run.font.bold = True
        run.font.size = Pt(14)
        para.space_before = Pt(0)
        para.space_after = Pt(1.67)

    def format_section_header(self, text: str):
        para = self.output_doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text.upper())
        run.font.bold = True
        run.font.size = Pt(12)
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(10)
        
        # Add border to the same paragraph
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')
        pBdr.append(bottom)
        para._element.get_or_add_pPr().append(pBdr)

    def format_education(self, education: dict):
        self.format_section_header("EDUCATION")
        
        uni_para = self.output_doc.add_paragraph()
        uni_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        uni_run = uni_para.add_run(education['university'])
        uni_run.font.bold = True
        uni_run.font.size = Pt(11)
        uni_para.space_before = Pt(0)
        uni_para.space_after = Pt(0.67)
        
        degree_para = self.output_doc.add_paragraph()
        degree_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        degree_run = degree_para.add_run(education['degree'])
        degree_run.font.size = Pt(11)
        degree_para.space_before = Pt(0)
        degree_para.space_after = Pt(1.67)

    def format_technical_skills(self, skills: dict):
        self.format_section_header("TECHNICAL SKILLS")
        
        # Create two-column layout for skills with minimal spacing
        table = self.output_doc.add_table(rows=1, cols=2)
        table.autofit = True
        table.allow_autofit = True
        table.style = 'Table Grid'
        table.style.paragraph_format.space_before = Pt(0)
        table.style.paragraph_format.space_after = Pt(0)
        table.style.paragraph_format.line_spacing = 1.0
        left_cell, right_cell = table.rows[0].cells
        
        # Split skills into two columns dynamically based on categories
        skills_items = list(skills.items())
        mid_point = len(skills_items) // 2
        
        # Left column
        for category, items in skills_items[:mid_point]:
            if items:
                p = left_cell.add_paragraph()
                category_run = p.add_run(f"• {category}: ")
                category_run.font.bold = True
                category_run.font.size = Pt(11)
                skills_run = p.add_run(', '.join(items))
                skills_run.font.size = Pt(11)
        
        # Right column
        for category, items in skills_items[mid_point:]:
            if items:
                p = right_cell.add_paragraph()
                category_run = p.add_run(f"• {category}: ")
                category_run.font.bold = True
                category_run.font.size = Pt(11)
                skills_run = p.add_run(', '.join(items))
                skills_run.font.size = Pt(11)

    def format_experience(self, experience: list):
        self.format_section_header("PROFESSIONAL EXPERIENCE")
        
        for job in experience:
            comp_para = self.output_doc.add_paragraph()
            comp_para.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
            
            comp_run = comp_para.add_run(job['company'])
            comp_run.font.bold = True
            comp_run.font.size = Pt(11)
            
            comp_para.add_run('\t')
            date_run = comp_para.add_run(job['dates'])
            date_run.font.bold = True
            date_run.font.size = Pt(11)
            comp_para.space_before = Pt(0)
            comp_para.space_after = Pt(0.67)
            
            title_para = self.output_doc.add_paragraph()
            title_run = title_para.add_run(job['title'])
            title_run.font.bold = True
            title_run.font.italic = True
            title_run.font.size = Pt(11)
            title_para.space_before = Pt(0)
            title_para.space_after = Pt(0.67)
            
            for bullet in job['bullets']:
                bullet_para = self.output_doc.add_paragraph()
                bullet_para.paragraph_format.left_indent = Inches(0.15)
                bullet_para.paragraph_format.first_line_indent = Inches(-0.15)
                bullet_run = bullet_para.add_run(f"• {bullet}")
                bullet_run.font.size = Pt(11)
                bullet_para.space_before = Pt(0)
                bullet_para.space_after = Pt(0.67)
            
            # Add space after each job section
            space_para = self.output_doc.add_paragraph()
            space_para.space_after = Pt(12)

    def set_margins(self):
        section = self.output_doc.sections[0]
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        
        # Add document-level spacing controls
        style = self.output_doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0  # Single line spacing
        
        # Adjust paragraph format for the entire document
        for paragraph in self.output_doc.paragraphs:
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    def create_formatted_resume(self, content: dict):
        self.set_margins()
        self.add_header_image()
        self.format_name(content['name'])
        self.format_education(content['education'])
        self.format_technical_skills(content['technical_skills'])
        self.format_experience(content['experience'])

def sanitize_text(text):
    """
    Removes or replaces characters that are invalid in XML.
    """
    # Remove NULL bytes
    text = text.replace('\x00', '')
    
    # Remove control characters except for common whitespace
    text = re.sub(r'[\x01-\x1F\x7F]', '', text)
    
    return text

def extract_text_from_pdf(pdf_file):
    try:
        # Extract text from PDF
        text = extract_text(pdf_file)
        if not text.strip():
            raise ValueError("PDF extraction resulted in empty text")
        
        # Sanitize text to remove invalid XML characters
        sanitized_text = sanitize_text(text)
        
        # Create a Document with the sanitized text
        doc = Document()
        for line in sanitized_text.split('\n'):
            doc.add_paragraph(line)
        
        return doc

    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def process_resume(uploaded_file):
    try:
        file_copy = io.BytesIO(uploaded_file.read())
        uploaded_file.seek(0)
        
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            doc = extract_text_from_pdf(file_copy)
        elif file_extension == 'docx':
            doc = Document(file_copy)
        else:
            return None, f"Unsupported file format: {file_extension}"
        
        text_content = "\n".join([para.text for para in doc.paragraphs])
        
        processor = ResumeProcessor(doc)
        content = processor.extract_content(text_content)
        processor.create_formatted_resume(content)
        
        return processor.output_doc, "Success"
        
    except Exception as e:
        logger.error(f"Error processing resume: {str(e)}")
        return None, f"Error: {str(e)}"

def create_streamlit_interface():
    st.image('optomi_logo.png', width=400)
    st.title('Resume Formatter')
    st.write("Upload a resume (PDF or DOCX) to standardize formatting.")
    
    uploaded_file = st.file_uploader("Upload Resume", type=['docx', 'pdf'])
    
    if uploaded_file:
        with st.spinner('Processing document...'):
            doc, message = process_resume(uploaded_file)
            
            if doc:
                bio = io.BytesIO()
                doc.save(bio)
                
                st.success("Document processed successfully!")
                
                original_name = uploaded_file.name
                output_filename = f"formatted_{original_name.rsplit('.', 1)[0]}.docx"
                
                st.download_button(
                    label="Download Formatted Resume",
                    data=bio.getvalue(),
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.error(message)

if __name__ == "__main__":
    create_streamlit_interface()
